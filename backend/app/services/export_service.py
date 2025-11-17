"""
Service xuất đề cương học phần ra file Word (.docx)
Sử dụng python-docx
"""
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from typing import List, Optional, Dict, Any
from app.models import Course, Program, CLO, Assessment, Question, CoursePrerequisite, PrerequisiteType, ConditionType, Rubric

def _replace_text_in_paragraph(paragraph, replacements):
    for key, value in replacements.items():
        if key in paragraph.text:
            inline = paragraph.runs
            for run in inline:
                if key in run.text:
                    run.text = run.text.replace(key, value)

def _replace_text_in_document(doc: Document, replacements: Dict[str, str]):
    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_text_in_paragraph(paragraph, replacements)

def _build_template_replacements(
    course: Course,
    program: Optional[Program],
    clos: List[CLO],
    assessments: List[Assessment],
    question_data: List[Dict[str, Any]],
    prerequisites: List[CoursePrerequisite],
    rubrics: List[Rubric],
    instructor_name: str,
    instructor_email: str,
    instructor_title: str,
    department: str,
    academic_year: str,
    include_prereqs: bool,
    include_rubrics: bool,
    session
) -> Dict[str, str]:
    replacements: Dict[str, str] = {
        "{{COURSE_CODE}}": course.code or "",
        "{{COURSE_NAME}}": course.title or "",
        "{{COURSE_CREDITS}}": str(course.credits or ""),
        "{{COURSE_DESCRIPTION}}": course.description or "(Chưa có mô tả)",
        "{{PROGRAM_NAME}}": program.name if program else "",
        "{{INSTRUCTOR_NAME}}": instructor_name,
        "{{INSTRUCTOR_EMAIL}}": instructor_email,
        "{{INSTRUCTOR_TITLE}}": instructor_title or "",
        "{{DEPARTMENT}}": department or "",
        "{{ACADEMIC_YEAR}}": academic_year or "",
    }
    
    clos_text = []
    for clo in clos:
        clos_text.append(f"{clo.code}: {clo.verb} {clo.text} (Bloom: {clo.bloom_level})")
    replacements["{{CLO_LIST}}"] = "\n".join(clos_text) if clos_text else "(Chưa có CLO)"
    
    if include_prereqs:
        prereq_text = []
        for prereq in prerequisites:
            prereq_course = session.get(Course, prereq.prereq_course_id)
            if prereq_course:
                prereq_text.append(format_condition(prereq, prereq_course))
        replacements["{{PREREQ_LIST}}"] = "\n".join(prereq_text) if prereq_text else "(Không có điều kiện tiên quyết)"
    else:
        replacements["{{PREREQ_LIST}}"] = "(Không bao gồm điều kiện tiên quyết)"
    
    assessment_lines = []
    for assessment in assessments:
        questions_for_assessment = [
            qd for qd in question_data 
            if qd["assessment"].id == assessment.id
        ]
        clo_mappings = []
        for qd in questions_for_assessment:
            for question in qd["questions"]:
                if question.clo_ids:
                    clo_mappings.extend([f"Q{question.id}→CLO{cid}" for cid in question.clo_ids])
        mapping_text = ", ".join(clo_mappings) if clo_mappings else "(Chưa map)"
        assessment_lines.append(
            f"{assessment.code} - {assessment.title} ({assessment.weight * 100:.1f}%) : {mapping_text}"
        )
    replacements["{{ASSESSMENT_LIST}}"] = "\n".join(assessment_lines) if assessment_lines else "(Chưa có kế hoạch đánh giá)"
    
    if include_rubrics and rubrics:
        rubric_lines = []
        for rubric in rubrics:
            rubric_lines.append(f"{rubric.name}: {rubric.description or ''}")
        replacements["{{RUBRIC_LIST}}"] = "\n".join(rubric_lines) if rubric_lines else "(Chưa có rubric)"
    elif include_rubrics:
        replacements["{{RUBRIC_LIST}}"] = "(Chưa có rubric)"
    else:
        replacements["{{RUBRIC_LIST}}"] = "(Không bao gồm rubric)"
    
    return replacements

def _generate_from_template(template_bytes: bytes, replacements: Dict[str, str]) -> bytes:
    doc = Document(io.BytesIO(template_bytes))
    _replace_text_in_document(doc, replacements)
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()

def format_condition(prereq: CoursePrerequisite, prereq_course: Course) -> str:
    """Format điều kiện tiên quyết thành text tiếng Việt"""
    condition_type = prereq.condition_type
    payload = prereq.condition_payload or {}
    
    course_name = f"{prereq_course.title} ({prereq_course.code})"
    
    if condition_type == ConditionType.PASS_COURSE:
        return f"Đã hoàn thành môn học {course_name}"
    elif condition_type == ConditionType.CLO_ACHIEVEMENT:
        required_ratio = payload.get("required_ratio", 0.66)
        return f"Đạt {int(required_ratio * 100)}% CLOs của môn học {course_name}"
    elif condition_type == ConditionType.PLO_THRESHOLD:
        threshold = payload.get("threshold", 0.6)
        return f"Đạt PLO threshold {threshold} của môn học {course_name}"
    elif condition_type == ConditionType.MIN_SCORE:
        min_score = payload.get("min_score", 5.0)
        return f"Điểm tối thiểu {min_score} trong môn học {course_name}"
    
    return f"Điều kiện: {course_name}"

def generate_course_docx(
    session,
    course: Course,
    program: Optional[Program],
    clos: List[CLO],
    assessments: List[Assessment],
    question_data: List[Dict[str, Any]],
    prerequisites: List[CoursePrerequisite],
    rubrics: List[Rubric],
    instructor_name: str,
    instructor_email: str,
    instructor_title: str,
    department: str,
    academic_year: str,
    include_prereqs: bool = True,
    include_rubrics: bool = True,
    template_bytes: Optional[bytes] = None
) -> bytes:
    """
    Tạo file Word đề cương học phần
    
    Structure:
    1. Trang bìa
    2. Mục lục
    3. Mô tả học phần
    4. Mục tiêu học phần (CLOs)
    5. Điều kiện tiên quyết
    6. Ma trận CLO -> PLO
    7. Kế hoạch đánh giá
    8. Rubrics (nếu có)
    """
    if template_bytes:
        replacements = _build_template_replacements(
            course,
            program,
            clos,
            assessments,
            question_data,
            prerequisites,
            rubrics,
            instructor_name,
            instructor_email,
            instructor_title,
            department,
            academic_year,
            include_prereqs,
            include_rubrics,
            session
        )
        return _generate_from_template(template_bytes, replacements)
    
    doc = Document()
    
    # Setup font cho tiếng Việt
    # python-docx tự động hỗ trợ UTF-8
    
    # ========== TRANG BÌA ==========
    title_para = doc.add_heading('ĐỀ CƯƠNG HỌC PHẦN', 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # Thông tin chương trình
    if program:
        p = doc.add_paragraph()
        p.add_run('Chương trình đào tạo: ').bold = True
        p.add_run(program.name)
    
    # Thông tin môn học
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Mã học phần: ').bold = True
    p.add_run(course.code)
    
    p = doc.add_paragraph()
    p.add_run('Tên học phần: ').bold = True
    p.add_run(course.title)
    
    p = doc.add_paragraph()
    p.add_run('Số tín chỉ: ').bold = True
    p.add_run(str(course.credits))
    
    if academic_year:
        p = doc.add_paragraph()
        p.add_run('Năm học: ').bold = True
        p.add_run(academic_year)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Giảng viên: ').bold = True
    p.add_run(instructor_name)
    
    if instructor_title:
        p = doc.add_paragraph()
        p.add_run('Chức danh: ').bold = True
        p.add_run(instructor_title)
    
    p = doc.add_paragraph()
    p.add_run('Email: ').bold = True
    p.add_run(instructor_email)
    
    if department:
        p = doc.add_paragraph()
        p.add_run('Bộ môn: ').bold = True
        p.add_run(department)
    
    doc.add_page_break()
    
    # ========== MỤC LỤC ==========
    doc.add_heading('MỤC LỤC', 1)
    doc.add_paragraph('1. Mô tả học phần')
    doc.add_paragraph('2. Mục tiêu học phần (CLOs)')
    if include_prereqs and prerequisites:
        doc.add_paragraph('3. Điều kiện tiên quyết')
    doc.add_paragraph('4. Ma trận CLO-PLO')
    doc.add_paragraph('5. Kế hoạch đánh giá')
    if include_rubrics:
        doc.add_paragraph('6. Rubrics đánh giá')
    
    doc.add_page_break()
    
    # ========== MÔ TẢ HỌC PHẦN ==========
    doc.add_heading('1. MÔ TẢ HỌC PHẦN', 1)
    if course.description:
        doc.add_paragraph(course.description)
    else:
        doc.add_paragraph('(Chưa có mô tả)')
    
    # ========== MỤC TIÊU HỌC PHẦN (CLOs) ==========
    doc.add_heading('2. MỤC TIÊU HỌC PHẦN (CLOs)', 1)
    
    if clos:
        for i, clo in enumerate(clos, 1):
            p = doc.add_paragraph(style='List Number')
            p.add_run(f'{clo.code}: ').bold = True
            p.add_run(f'{clo.verb} {clo.text}')
            p.add_run(f' (Bloom: {clo.bloom_level})').italic = True
    else:
        doc.add_paragraph('(Chưa có CLO)')
    
    # ========== ĐIỀU KIỆN TIÊN QUYẾT ==========
    if include_prereqs and prerequisites:
        doc.add_heading('3. ĐIỀU KIỆN TIÊN QUYẾT', 1)
        
        # Nhóm theo type
        strict_prereqs = [p for p in prerequisites if p.type == PrerequisiteType.STRICT]
        coreq_prereqs = [p for p in prerequisites if p.type == PrerequisiteType.COREQ]
        recommended_prereqs = [p for p in prerequisites if p.type == PrerequisiteType.RECOMMENDED]
        
        if strict_prereqs:
            doc.add_heading('3.1. Môn học tiên quyết (Bắt buộc)', 2)
            for prereq in strict_prereqs:
                prereq_course = session.get(Course, prereq.prereq_course_id)
                if prereq_course:
                    condition_text = format_condition(prereq, prereq_course)
                    doc.add_paragraph(condition_text, style='List Bullet')
        
        if coreq_prereqs:
            doc.add_heading('3.2. Môn học đồng thời (Co-requisite)', 2)
            for prereq in coreq_prereqs:
                prereq_course = session.get(Course, prereq.prereq_course_id)
                if prereq_course:
                    condition_text = format_condition(prereq, prereq_course)
                    doc.add_paragraph(condition_text, style='List Bullet')
        
        if recommended_prereqs:
            doc.add_heading('3.3. Môn học khuyến nghị (Recommended)', 2)
            for prereq in recommended_prereqs:
                prereq_course = session.get(Course, prereq.prereq_course_id)
                if prereq_course:
                    condition_text = format_condition(prereq, prereq_course)
                    doc.add_paragraph(condition_text, style='List Bullet')
    elif include_prereqs:
        doc.add_heading('3. ĐIỀU KIỆN TIÊN QUYẾT', 1)
        doc.add_paragraph('(Không có điều kiện tiên quyết)')
    
    # ========== MA TRẬN CLO-PLO ==========
    doc.add_heading('4. MA TRẬN CLO-PLO', 1)
    
    # Lấy mappings
    from app.models import CLOPLOMapping, PLO
    from sqlmodel import select
    
    if clos:
        # Lấy PLOs của program
        plos = []
        if program:
            plos = session.exec(select(PLO).where(PLO.program_id == program.id)).all()
        
        if plos:
            # Lấy mappings
            clo_ids = [clo.id for clo in clos]
            plo_ids = [plo.id for plo in plos]
            statement = select(CLOPLOMapping).where(
                CLOPLOMapping.clo_id.in_(clo_ids),
                CLOPLOMapping.plo_id.in_(plo_ids)
            )
            mappings = session.exec(statement).all()
            mapping_dict = {(m.clo_id, m.plo_id): m.contribution_level for m in mappings}
            
            # Tạo bảng
            table = doc.add_table(rows=1, cols=len(plos) + 1)
            table.style = 'Light Grid Accent 1'
            
            # Header
            header_cells = table.rows[0].cells
            header_cells[0].text = 'CLO'
            for i, plo in enumerate(plos, 1):
                header_cells[i].text = plo.code
            
            # Rows
            for clo in clos:
                row_cells = table.add_row().cells
                row_cells[0].text = clo.code
                for i, plo in enumerate(plos, 1):
                    level = mapping_dict.get((clo.id, plo.id), '-')
                    row_cells[i].text = level
        else:
            doc.add_paragraph('(Chưa có PLO)')
    else:
        doc.add_paragraph('(Chưa có CLO)')
    
    # ========== KẾ HOẠCH ĐÁNH GIÁ ==========
    doc.add_heading('5. KẾ HOẠCH ĐÁNH GIÁ', 1)
    
    if assessments:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Mã đánh giá'
        header_cells[1].text = 'Tên đánh giá'
        header_cells[2].text = 'Trọng số (%)'
        header_cells[3].text = 'Câu hỏi → CLO'
        
        for assessment in assessments:
            row_cells = table.add_row().cells
            row_cells[0].text = assessment.code
            row_cells[1].text = assessment.title
            row_cells[2].text = f"{assessment.weight * 100:.1f}%"
            
            # Lấy questions của assessment này
            questions_for_assessment = [
                qd for qd in question_data 
                if qd["assessment"].id == assessment.id
            ]
            
            clo_mappings = []
            for qd in questions_for_assessment:
                for question in qd["questions"]:
                    if question.clo_ids:
                        clo_mappings.extend([f"Q{question.id}→CLO{cid}" for cid in question.clo_ids])
            
            row_cells[3].text = ", ".join(clo_mappings) if clo_mappings else "(Chưa map)"
    else:
        doc.add_paragraph('(Chưa có kế hoạch đánh giá)')
    
    # ========== RUBRICS ==========
    if include_rubrics:
        doc.add_heading('6. RUBRICS ĐÁNH GIÁ', 1)
        
        if rubrics:
            # Nhóm rubrics theo CLO
            rubric_by_clo = {}
            for rubric in rubrics:
                if rubric.clo_id:
                    if rubric.clo_id not in rubric_by_clo:
                        rubric_by_clo[rubric.clo_id] = []
                    rubric_by_clo[rubric.clo_id].append(rubric)
            
            # Xuất rubric cho từng CLO
            for clo in clos:
                if clo.id in rubric_by_clo:
                    rubric = rubric_by_clo[clo.id][0]  # Mỗi CLO chỉ có 1 rubric
                    
                    # Tiêu đề CLO
                    p = doc.add_paragraph()
                    p.add_run(f'{clo.code}: ').bold = True
                    p.add_run(f'{clo.verb} {clo.text}')
                    
                    # Tên rubric
                    p = doc.add_paragraph()
                    p.add_run(f'Rubric: {rubric.name}').bold = True
                    
                    if rubric.description:
                        doc.add_paragraph(rubric.description, style='Intense Quote')
                    
                    # Bảng tiêu chí
                    if rubric.criteria:
                        table = doc.add_table(rows=1, cols=5)
                        table.style = 'Light Grid Accent 1'
                        
                        # Header
                        header_cells = table.rows[0].cells
                        header_cells[0].text = 'Tiêu chí'
                        header_cells[1].text = 'Xuất sắc (4)'
                        header_cells[2].text = 'Tốt (3)'
                        header_cells[3].text = 'Đạt (2)'
                        header_cells[4].text = 'Chưa đạt (1)'
                        
                        # Dữ liệu
                        for criterion_key, criterion in rubric.criteria.items():
                            row_cells = table.add_row().cells
                            row_cells[0].text = criterion.get('name', '')
                            row_cells[1].text = criterion.get('levels', {}).get('4', '')
                            row_cells[2].text = criterion.get('levels', {}).get('3', '')
                            row_cells[3].text = criterion.get('levels', {}).get('2', '')
                            row_cells[4].text = criterion.get('levels', {}).get('1', '')
                    
                    doc.add_paragraph()  # Spacing
        else:
            doc.add_paragraph('(Chưa có rubric)')
    
    # ========== CHỮ KÝ ==========
    doc.add_page_break()
    doc.add_heading('XÁC NHẬN', 1)
    doc.add_paragraph()
    doc.add_paragraph('Giảng viên phụ trách')
    doc.add_paragraph()
    doc.add_paragraph('_________________________')
    doc.add_paragraph(instructor_name)
    
    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes.read()

